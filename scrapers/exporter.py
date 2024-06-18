import csv
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def export_to_txt(company: str, headers: list, questions: list) -> None:
    """Export interview questions to a TXT file"""
    try:
        filename = f"{company}_questions.txt"
        with open(filename, "w") as file:
            file.write(f"Interview Questions for {company}\n\n")
            for question in questions:
                for header in headers:
                    file.write(f"{header}: {question.get(header, '')}\n")
                file.write("\n")
    except Exception as e:
        raise Exception(f"An error occurred while exporting to TXT: {e}")


def export_to_docx(company: str, headers: list, questions: list) -> None:
    """Export interview questions to a DOCX file"""
    try:
        filename = f"{company}_questions.docx"
        doc = Document()
        doc.add_heading(f"Interview Questions for {company}\n", level=1)
        for question in questions:
            for header in headers:
                doc.add_paragraph(
                    f"{header}: {question.get(header, '')}", style="List Bullet"
                )
            doc.add_paragraph()
        doc.save(filename)
    except Exception as e:
        raise Exception(f"An error occurred while exporting to DOCX: {e}")


def export_to_csv(company: str, headers: list, questions: list) -> None:
    """Export interview questions to a CSV file"""
    filename = f"{company}_questions.csv"
    try:
        with open(filename, mode="w", newline="", encoding="utf-8") as file:
            writer = csv.DictWriter(file, fieldnames=headers)
            writer.writeheader()
            for question in questions:
                writer.writerow(
                    {header: question.get(header, "") for header in headers}
                )
    except IOError as e:
        raise Exception(f"Failed to export interview questions: {e}")


def export_to_pdf(company: str, headers: list, questions: list) -> None:
    """Export questions to a PDF file"""
    pdf_filename = f"{company}_interview_questions.pdf"
    c = canvas.Canvas(pdf_filename, pagesize=letter)
    c.setFont("Helvetica", 12)

    c.setFont("Helvetica-Bold", 14)
    header = f"Interview Questions for {company}"
    c.drawString(100, 750, header)
    c.setFont("Helvetica", 12)
    y -= 20

    y = 730
    for question in questions:
        for header in headers:
            question_str = f"{header}: {question.get(header, '')}"
            c.drawString(100, y, question_str)
            y -= 20
        y -= 20
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = 750

    c.save()
